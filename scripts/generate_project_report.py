"""Format the user-supplied Tent House report source into an editable DOCX."""
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"C:\Users\ASUS\.codex\attachments\c7217399-8205-46db-a6c2-9e81c0a9ab7b\pasted-text.txt")
OUTPUT = ROOT / "docs" / "MERN_Tent_House_Project_Report.docx"

MAJOR_HEADINGS = {
    "BONAFIDE CERTIFICATE", "DECLARATION", "ACKNOWLEDGEMENT", "ABSTRACT",
    "TABLE OF CONTENTS", "LIST OF TABLES", "LIST OF FIGURES / SCREENSHOTS",
    "REFERENCES", "APPENDIX A: SOURCE CODE SNIPPETS",
    "CHAPTER 1: INTRODUCTION", "CHAPTER 2: LITERATURE SURVEY & FEASIBILITY ANALYSIS",
    "CHAPTER 3: SYSTEM REQUIREMENT SPECIFICATIONS (SRS)",
    "CHAPTER 4: SYSTEM DESIGN & ARCHITECTURE",
    "CHAPTER 5: IMPLEMENTATION & MODULE DETAILS",
    "CHAPTER 6: SOFTWARE TESTING & QUALITY ASSURANCE",
    "CHAPTER 7: SYSTEM SCREENSHOTS & UI VISUALIZATIONS",
    "CHAPTER 8: DEPLOYMENT ARCHITECTURE & OPERATIONS",
    "CHAPTER 9: CONCLUSION & FUTURE ENHANCEMENTS",
}

def configure(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.right_margin = Inches(1)
    sec.left_margin = Inches(1.25)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer.add_run()._r.append(field)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

def paragraph(doc, value, center=False, bold=False, size=12, mono=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(value)
    r.font.name = "Courier New" if mono else "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    return p

def clean_source(source):
    # The pasted source uses literal backslashes as line separators.
    source = source.replace("\\r\\n", "\n").replace("\\n", "\n")
    source = source.replace("PAGE BREAK", "\n[[PAGE_BREAK]]\n")
    source = source.replace("\\", "\n")
    # Do not put secret-looking values into a report, even when an example was pasted.
    source = re.sub(r"(?m)^(MONGO_URI|MONGODB_URI|JWT_SECRET|CLOUDINARY_API_KEY|CLOUDINARY_API_SECRET)=.*$", r"\1=<REDACTED – set in hosting environment>", source)
    return source

def insert_readability_breaks(value):
    headings = [
        r"CHAPTER [1-9]:", r"\d+\.\d+(?:\.\d+)?:", r"Table \d+\.\d+:",
        r"Figure \d+\.\d+:", r"A\.\d+ ", r"Code snippet", r"Plaintext",
        r"REFERENCES", r"APPENDIX A:", r"Prerequisites", r"Installation and Execution Steps",
    ]
    return re.sub("(" + "|".join(headings) + ")", r"\n\1", value)

def is_heading(line):
    plain = line.strip()
    return (
        plain in MAJOR_HEADINGS
        or bool(re.match(r"^(CHAPTER [1-9]:|\d+\.\d+(?:\.\d+)?\s|\d+\.\d+(?:\.\d+)?:)", plain))
        or plain.startswith(("Table ", "Figure ", "APPENDIX "))
    )

def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    paragraph(doc, "A PROJECT REPORT ON", center=True, bold=True, size=16)
    paragraph(doc, "MUMTAZ DAAWAT DECOR & TENT", center=True, bold=True, size=20)
    paragraph(doc, "A Full-Stack MERN Event Booking and Tent-House Management Platform", center=True, bold=True, size=14)
    paragraph(doc, "Submitted in partial fulfillment of the requirements for the award of the degree of", center=True)
    paragraph(doc, "BACHELOR OF TECHNOLOGY IN COMPUTER SCIENCE AND ENGINEERING", center=True, bold=True, size=15)
    paragraph(doc, "Submitted By:", center=True)
    paragraph(doc, "Khushnudi Mumtaz", center=True, bold=True, size=15)
    paragraph(doc, "Registration No.: 12500653", center=True)
    paragraph(doc, "Under the Guidance of: [Supervisor Name]", center=True)
    paragraph(doc, "DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING", center=True, bold=True)
    paragraph(doc, "[NAME OF YOUR UNIVERSITY / INSTITUTE]", center=True, bold=True)
    paragraph(doc, "ACADEMIC YEAR 2025–2026", center=True)
    doc.add_page_break()

def add_certificate_page(doc):
    paragraph(doc, "TRAINING CERTIFICATE FROM ORGANIZATION", center=True, bold=True, size=14)
    for _ in range(17):
        doc.add_paragraph()
    paragraph(doc, "INTENTIONALLY LEFT BLANK FOR OFFICIAL TRAINING / INTERNSHIP CERTIFICATE ATTACHMENT", center=True, size=11)
    doc.add_page_break()

def add_formatted_page(doc, block):
    block = insert_readability_breaks(block)
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    mono_mode = False
    for line in lines:
        if line in {"Plaintext", "Code snippet", "JavaScript", "Bash"}:
            mono_mode = True
            paragraph(doc, line, bold=True, size=10, mono=True)
            continue
        if is_heading(line):
            mono_mode = False
            paragraph(doc, line, center=line in MAJOR_HEADINGS or line.startswith("CHAPTER"), bold=True, size=14 if line in MAJOR_HEADINGS or line.startswith("CHAPTER") else 12)
            continue
        looks_like_diagram = line.startswith(("+", "|", "#", "- ", "[ ")) or "---->" in line
        paragraph(doc, line, mono=mono_mode or looks_like_diagram, size=9 if mono_mode or looks_like_diagram else 12)

def main():
    raw = SOURCE.read_text(encoding="utf-8")
    raw = clean_source(raw)
    blocks = [item.strip() for item in raw.split("[[PAGE_BREAK]]") if item.strip()]

    doc = Document()
    configure(doc)
    add_cover(doc)

    certificate_added = False
    for index, block in enumerate(blocks):
        # The supplied first block is cover content already formatted above.
        if index == 0:
            continue
        if "DEDICATED TRAINING CERTIFICATE PAGE" in block:
            add_certificate_page(doc)
            certificate_added = True
            continue
        add_formatted_page(doc, block)
        doc.add_page_break()

    if not certificate_added:
        add_certificate_page(doc)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")

if __name__ == "__main__":
    main()
